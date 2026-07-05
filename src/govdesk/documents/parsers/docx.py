# SPDX-FileCopyrightText: 2026 GovDesk-Mitwirkende
#
# SPDX-License-Identifier: EUPL-1.2

"""DOCX-Parser (python-docx): Überschriften-Styles werden Struktur."""

import io
import re

from docx import Document as DocxDocument

from govdesk.documents.parsers.base import Block, ParsedDocument

_HEADING_STYLE = re.compile(r"^(?:heading|überschrift)\s*(\d)$", re.IGNORECASE)


class DocxParser:
    def parse(self, data: bytes) -> ParsedDocument:
        docx = DocxDocument(io.BytesIO(data))
        blocks: list[Block] = []

        for paragraph in docx.paragraphs:
            text = " ".join(paragraph.text.split())
            if not text:
                continue
            style_name = (paragraph.style.name or "") if paragraph.style else ""
            match = _HEADING_STYLE.match(style_name.strip())
            if match:
                blocks.append(Block(text=text, heading_level=int(match.group(1))))
            elif style_name.lower() in ("title", "titel"):
                blocks.append(Block(text=text, heading_level=1))
            else:
                blocks.append(Block(text=text))

        # Tabellen zeilenweise als Text übernehmen
        for table in docx.tables:
            for row in table.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    blocks.append(Block(text=line))

        return ParsedDocument(blocks=blocks)
