# SPDX-FileCopyrightText: 2026 GovDesk-Mitwirkende
#
# SPDX-License-Identifier: EUPL-1.2

"""Export eines Editor-Dokuments (HTML) nach DOCX und ODF.

Das gespeicherte HTML nutzt nur eine kleine Tag-Menge (Überschriften, Absätze,
Listen, Fett/Kursiv/Unterstrichen). Wir parsen es mit der Stdlib und bauen daraus
saubere Office-Dokumente (python-docx / odfpy) — beides weiterverarbeitbar.
"""

import io
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser

_BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "li", "blockquote"}
_BOLD = {"b", "strong"}
_ITALIC = {"i", "em"}
_UNDERLINE = {"u"}


@dataclass
class _Run:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False


@dataclass
class _Block:
    type: str  # "h1".."h4" | "p" | "li" | "quote"
    ordered: bool = False
    runs: list[_Run] = field(default_factory=list)


class _BlockParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[_Block] = []
        self._cur: _Block | None = None
        self._b = self._i = self._u = 0
        self._lists: list[str] = []

    def _flush(self) -> None:
        if self._cur and any(r.text.strip() for r in self._cur.runs):
            self.blocks.append(self._cur)
        self._cur = None

    def handle_starttag(self, tag: str, attrs: list) -> None:
        if tag in ("ul", "ol"):
            self._lists.append(tag)
            return
        if tag == "br" and self._cur is not None:
            self._cur.runs.append(_Run("\n"))
            return
        if tag in _BLOCK_TAGS:
            self._flush()
            if tag == "li":
                ordered = bool(self._lists) and self._lists[-1] == "ol"
                self._cur = _Block("li", ordered=ordered)
            elif tag == "blockquote":
                self._cur = _Block("quote")
            else:
                self._cur = _Block(tag if tag[0] == "h" else "p")
            return
        if tag in _BOLD:
            self._b += 1
        elif tag in _ITALIC:
            self._i += 1
        elif tag in _UNDERLINE:
            self._u += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in ("ul", "ol"):
            if self._lists:
                self._lists.pop()
        elif tag in _BLOCK_TAGS:
            self._flush()
        elif tag in _BOLD and self._b:
            self._b -= 1
        elif tag in _ITALIC and self._i:
            self._i -= 1
        elif tag in _UNDERLINE and self._u:
            self._u -= 1

    def handle_data(self, data: str) -> None:
        text = re.sub(r"\s+", " ", data)
        if not text:
            return
        if self._cur is None:
            self._cur = _Block("p")
        self._cur.runs.append(_Run(text, bool(self._b), bool(self._i), bool(self._u)))

    def close(self) -> None:
        super().close()
        self._flush()


def _parse(html: str) -> list[_Block]:
    parser = _BlockParser()
    parser.feed(html or "")
    parser.close()
    return parser.blocks


def to_docx(title: str, html: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)
    for block in _parse(html):
        if block.type in ("h1", "h2", "h3", "h4"):
            para = doc.add_paragraph(style=f"Heading {block.type[1]}")
        elif block.type == "li":
            para = doc.add_paragraph(style="List Number" if block.ordered else "List Bullet")
        elif block.type == "quote":
            para = doc.add_paragraph(style="Quote")
        else:
            para = doc.add_paragraph()
        for run in block.runs:
            r = para.add_run(run.text)
            r.bold = run.bold
            r.italic = run.italic
            r.underline = run.underline
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _runs_markup(runs: list[_Run]) -> str:
    """Runs → reportlab-Inline-Markup (<b>/<i>/<u>), Text escaped."""
    from xml.sax.saxutils import escape

    parts = []
    for run in runs:
        text = escape(run.text)
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.underline:
            text = f"<u>{text}</u>"
        parts.append(text)
    return "".join(parts) or "&nbsp;"


def to_pdf(title: str, html: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, title=title,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    heading = {"h1": "Heading1", "h2": "Heading2", "h3": "Heading3", "h4": "Heading4"}
    from xml.sax.saxutils import escape as _esc

    flow = [Paragraph(_esc(title), styles["Title"]), Spacer(1, 0.4 * cm)]
    pending: list = []

    def flush_list() -> None:
        nonlocal pending
        if pending:
            ordered = pending[0][1]
            items = [ListItem(Paragraph(m, styles["BodyText"])) for m, _ in pending]
            flow.append(ListFlowable(items, bulletType="1" if ordered else "bullet"))
            pending = []

    for block in _parse(html):
        if block.type == "li":
            pending.append((_runs_markup(block.runs), block.ordered))
            continue
        flush_list()
        markup = _runs_markup(block.runs)
        style = styles[heading[block.type]] if block.type in heading else styles["BodyText"]
        flow.append(Paragraph(markup, style))
    flush_list()

    doc.build(flow)
    return buffer.getvalue()


def to_odf(title: str, html: str) -> bytes:
    from odf.opendocument import OpenDocumentText
    from odf.style import Style, TextProperties
    from odf.text import H, List, ListItem, P, Span

    doc = OpenDocumentText()
    styles = {}
    for name, props in (
        ("gd_bold", {"fontweight": "bold"}),
        ("gd_italic", {"fontstyle": "italic"}),
        ("gd_ul", {"textunderlinestyle": "solid", "textunderlinewidth": "auto"}),
    ):
        st = Style(name=name, family="text")
        st.addElement(TextProperties(**props))
        doc.automaticstyles.addElement(st)
        styles[name] = st

    def fill(element, runs: list[_Run]) -> None:
        for run in runs:
            style = None
            if run.bold:
                style = styles["gd_bold"]
            elif run.italic:
                style = styles["gd_italic"]
            elif run.underline:
                style = styles["gd_ul"]
            if style is not None:
                element.addElement(Span(stylename=style, text=run.text))
            else:
                element.addText(run.text)

    doc.text.addElement(H(outlinelevel=1, text=title))
    pending: List | None = None
    for block in _parse(html):
        if block.type == "li":
            if pending is None:
                pending = List()
                doc.text.addElement(pending)
            item = ListItem()
            para = P()
            fill(para, block.runs)
            item.addElement(para)
            pending.addElement(item)
            continue
        pending = None
        if block.type in ("h1", "h2", "h3", "h4"):
            element = H(outlinelevel=int(block.type[1]) + 1)
        else:
            element = P()
        fill(element, block.runs)
        doc.text.addElement(element)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
